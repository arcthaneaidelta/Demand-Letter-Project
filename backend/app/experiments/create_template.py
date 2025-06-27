from docx import Document

# Create a new Word document
doc = Document()

# Add a title
title = doc.add_heading('Contract Agreement', 0)

# Add client information section with placeholders
doc.add_heading('Client Information', level=1)
para = doc.add_paragraph('Client Name: ')
para.add_run('{{ client_name }}').bold = True
para = doc.add_paragraph('Address: ')
para.add_run('{{ address }}').bold = True
para = doc.add_paragraph('Service Date: ')
para.add_run('{{ service_date }}').bold = True

# Add services section
doc.add_heading('Services Provided', level=1)
doc.add_paragraph('The following services were provided:')

# Add a table for items
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Service'
hdr_cells[1].text = 'Price'

# Add table row template for items
row_cells = table.add_row().cells
row_cells[0].text = '{{ items[0].name }}'
row_cells[1].text = '{{ items[0].price }}'

row_cells = table.add_row().cells
row_cells[0].text = '{{ items[1].name }}'
row_cells[1].text = '{{ items[1].price }}'

# Add contact information
doc.add_heading('Contact Information', level=1)
para = doc.add_paragraph('For any inquiries, please contact us at: ')
para.add_run('{{ contact_email }}').italic = True

# Add company information
doc.add_paragraph()
para = doc.add_paragraph('Best regards,')
para = doc.add_paragraph('{{ company_name }}')

# Save the document
doc.save('./test_doc.docx')

print("Template created successfully: test_doc.docx") 